"""
engine/knowledge_render.py
تولید فایل‌های PDF و DOCX از مدل گزارش DANA (خروجی engine/knowledge_draft.build_report).

سازگار با اسکریپت‌های مرجع مهارت organizational-knowledge-skill
(scripts/make_dana_pdf.py و scripts/make_dana_docx.py) — با این تفاوت که این‌جا
دیتا-درایو و قابل‌استفاده در ربات است:
    - render_dana_docx(report, out_path) — همیشه کار می‌کند (فونت در خود فایل نام دارد)
    - render_dana_pdf(report, out_path) — به فونت عربی روی سرور نیاز دارد؛
      اگر پیدا نشود False برمی‌گرداند و فقط DOCX ارسال می‌شود.
"""

from __future__ import annotations

import logging
import os

logger = logging.getLogger(__name__)

try:
    import arabic_reshaper
    from bidi.algorithm import get_display

    _RTL_AVAILABLE = True
except Exception:  # pragma: no cover
    arabic_reshaper = None  # type: ignore
    get_display = None  # type: ignore
    _RTL_AVAILABLE = False


# ══════════════════════════════════════════════════════════════════════════════
# کشف فونت عربی (برای PDF)
# ══════════════════════════════════════════════════════════════════════════════

# مسیرهای شناخته‌شده روی ویندوز/لینوکس
_FONT_CANDIDATES = [
    # ویندوز
    r"C:\Windows\Fonts\tahoma.ttf",
    r"C:\Windows\Fonts\tahomabd.ttf",
    r"C:\Windows\Fonts\arial.ttf",
    # لینوکس (msttcorefonts, fonts-arabeyes و امثال آن)
    "/usr/share/fonts/truetype/msttcorefonts/Tahoma.ttf",
    "/usr/share/fonts/truetype/msttcorefonts/arial.ttf",
    "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
]

# برای فونت‌های عربی (با پوشش گلیف‌های ارائهٔ شکل‌دار) جستجو در دایرکتوری‌های رایج
_FONT_SCAN_DIRS = [
    "/usr/share/fonts/truetype",
    "/usr/local/share/fonts",
]


def _candidate_is_arabic(font_path: str) -> bool:
    """بر اساس نام فایل، حدس ساده‌ای از عربی‌بودن فونت می‌زند."""
    name = os.path.basename(font_path).lower()
    arabic_hints = ("tahoma", "arabic", "amiri", "noto.naskh", "naskh", "scheherazade", "vazir", "iran")
    return any(hint in name for hint in arabic_hints)


def _find_arabic_font() -> str | None:
    """اولین فونت موجود با پوشش عربی را پیدا می‌کند؛ در غیر این صورت None."""
    # اول مسیرهای شناخته‌شدهٔ ویندوز/لینوکس
    for path in _FONT_CANDIDATES:
        if os.path.exists(path):
            return path
    # بعد، اسکن سبک دایرکتوری‌های رایج فونت‌های لینوکس
    for base in _FONT_SCAN_DIRS:
        if not os.path.isdir(base):
            continue
        for root, _, files in os.walk(base):
            for fn in files:
                if not fn.lower().endswith((".ttf", ".otf")):
                    continue
                full = os.path.join(root, fn)
                if _candidate_is_arabic(full):
                    return full
    return None


# ══════════════════════════════════════════════════════════════════════════════
# PDF — مبتنی بر reportlab (سازگار با make_dana_pdf.py مهارت)
# ══════════════════════════════════════════════════════════════════════════════


def _fa(text: str) -> str:
    """بازشکل‌دهی و راست‌چین‌کردن متن فارسی برای نمایش در PDF."""
    if not _RTL_AVAILABLE:
        return text
    return get_display(arabic_reshaper.reshape(str(text)))


def render_dana_pdf(report: dict, out_path: str) -> bool:
    """
    مدل گزارش را به PDF تبدیل می‌کند.

    خروجی:
        True  — فایل ساخته شد
        False — فونت عربی روی سرور یافت نشد (فایل ساخته نشد)
    """
    if not _RTL_AVAILABLE:
        logger.warning("arabic_reshaper/python-bidi نصب نیست — ساخت PDF ممکن نیست.")
        return False

    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_RIGHT
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.platypus import (
            HRFlowable,
            Paragraph,
            SimpleDocTemplate,
            Spacer,
        )
    except ImportError:
        logger.warning("reportlab نصب نیست — ساخت PDF ممکن نیست.")
        return False

    font_path = _find_arabic_font()
    if font_path is None:
        logger.warning("فونت عربی روی سرور یافت نشد — فقط DOCX ارسال می‌شود.")
        return False

    os.makedirs(os.path.dirname(os.path.abspath(out_path)), exist_ok=True)

    # ثبت فونت: همین یک خانواده برای متن و عنوان (bold مصنوعی نیست، پس حالت ساده)
    pdfmetrics.registerFont(TTFont("KN", font_path))

    title_style = ParagraphStyle(
        "KNTitle", fontName="KN", fontSize=16, leading=22,
        alignment=TA_RIGHT, textColor=colors.HexColor("#1F4E79"), spaceAfter=6,
    )
    section_style = ParagraphStyle(
        "KNSection", fontName="KN", fontSize=12, leading=16,
        alignment=TA_RIGHT, textColor=colors.HexColor("#1F4E79"),
        spaceBefore=12, spaceAfter=4,
    )
    body_style = ParagraphStyle(
        "KNBody", fontName="KN", fontSize=10.5, leading=17,
        alignment=TA_RIGHT, textColor=colors.black,
    )
    meta_style = ParagraphStyle(
        "KNMeta", fontName="KN", fontSize=10, leading=16,
        alignment=TA_RIGHT, textColor=colors.HexColor("#444444"),
    )
    small_style = ParagraphStyle(
        "KNSmall", fontName="KN", fontSize=8.5, leading=13,
        alignment=TA_RIGHT, textColor=colors.HexColor("#666666"),
    )

    doc = SimpleDocTemplate(
        out_path, pagesize=A4,
        leftMargin=20 * mm, rightMargin=20 * mm,
        topMargin=15 * mm, bottomMargin=15 * mm,
        title=f"پیش‌نویس ثبت دانش در DANA — {report.get('title', '')}",
        author="WelderBot",
    )

    story: list = []

    # ── اطلاعات ثبت ────────────────────────────────────────────────────────
    story.append(Paragraph(_fa(f"📄 پیش‌نویس ثبت دانش در DANA"), title_style))
    story.append(HRFlowable(width="100%", thickness=1.2, color=colors.HexColor("#1F4E79")))
    story.append(Paragraph(_fa(f"نوع دانش: {report['type_label']}"), body_style))
    story.append(Paragraph(_fa(f"وضعیت QA: {report['qa_status']}"), body_style))
    story.append(Paragraph(_fa(f"بازبینی اپراتور: {report['operator_review']}"), body_style))

    # ── محتوا ──────────────────────────────────────────────────────────────
    story.append(Paragraph(_fa("────── محتوا ──────"), section_style))
    for label, value in report["content"]:
        story.append(Paragraph(_fa(f"<b>{label}</b>: {value}"), body_style))

    # ── فراداده ────────────────────────────────────────────────────────────
    story.append(Paragraph(_fa("────── فراداده ──────"), section_style))
    for label, value in report["metadata"]:
        story.append(Paragraph(_fa(f"• {label}: {value}"), meta_style))

    # ── منابع ──────────────────────────────────────────────────────────────
    story.append(Paragraph(_fa("────── منابع ──────"), section_style))
    story.append(Paragraph(_fa("پیوست‌ها:"), meta_style))
    for r in report["resources"]:
        story.append(Paragraph(_fa(f"  {r}"), meta_style))

    # ── موارد حل‌نشده ──────────────────────────────────────────────────────
    story.append(Paragraph(_fa("────── موارد حل‌نشده ──────"), section_style))
    for item in report["unresolved"]:
        story.append(Paragraph(_fa(f"• {item}"), meta_style))

    # ── چک‌لیست نهایی اپراتور ──────────────────────────────────────────────
    story.append(Paragraph(_fa("────── چک‌لیست نهایی اپراتور ──────"), section_style))
    for item in report["checklist"]:
        story.append(Paragraph(_fa(f"[ ] {item}"), body_style))

    story.append(Spacer(1, 10))
    story.append(HRFlowable(width="100%", thickness=0.8, color=colors.HexColor("#AAAAAA")))
    story.append(Paragraph(_fa(report["footer"]), small_style))

    doc.build(story)
    logger.info("PDF دانش ساخته شد: %s (فونت: %s)", out_path, font_path)
    return True


# ══════════════════════════════════════════════════════════════════════════════
# DOCX — مبتنی بر python-docx (سازگار با make_dana_docx.py مهارت)
# ══════════════════════════════════════════════════════════════════════════════

def _set_rtl(run) -> None:
    """ویژگی bidi راست‌به‌چپ را روی run فعال می‌کند (برای متن فارسی)."""
    try:
        run.font.rtl = True  # python-docx >= 1.1
    except Exception:
        try:
            from docx.oxml.ns import qn

            rPr = run._element.get_or_add_rPr()
            rtl = rPr.find(qn("w:rtl"))
            if rtl is None:
                rPr.append(rPr.makeelement(qn("w:rtl"), {}))
        except Exception:  # pragma: no cover
            pass


def render_dana_docx(report: dict, out_path: str) -> str:
    """
    مدل گزارش را به DOCX تبدیل می‌کند (همیشه بدون خطا، به‌شرط نصب python-docx).

    فونت‌ها فقط به اسم در فایل ذخیره می‌شوند (B Titr / B Nazanin) — سرور به آن‌ها
    نیاز ندارد؛ هنگام باز کردن در Word روی سیستمِ دارای فونت رندر می‌شوند.
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.section import WD_ORIENT
        from docx.oxml.ns import qn
        from docx.shared import Mm, Pt
    except ImportError:
        logger.warning("python-docx نصب نیست — ساخت DOCX ممکن نیست.")
        raise

    os.makedirs(os.path.dirname(os.path.abspath(out_path)), exist_ok=True)

    doc = Document()

    # A4 + حاشیه 18mm (مطابق make_dana_docx.py مهارت)
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Mm(18)
    section.right_margin = Mm(18)
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(18)

    # فونت پیش‌فرض سند
    normal = doc.styles["Normal"]
    normal.font.name = "B Nazanin"
    normal.font.size = Pt(11)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "B Nazanin")

    def _add_paragraph(text: str, *, size: float, bold: bool, color: str | None = None, space: int = 4):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(space)
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        run.font.name = "B Titr" if bold else "B Nazanin"
        if color:
            from docx.shared import RGBColor

            run.font.color.rgb = RGBColor.from_string(color)
        _set_rtl(run)
        return p

    def _add_heading(text: str):
        p = _add_paragraph(text, size=12, bold=True, color="1F4E79", space=4)
        return p

    # ── اطلاعات ثبت ────────────────────────────────────────────────────────
    _add_paragraph("پیش‌نویس ثبت دانش در DANA", size=17, bold=True, color="1F4E79", space=8)
    _add_paragraph(f"نوع دانش: {report['type_label']}", size=11, bold=False)
    _add_paragraph(f"وضعیت QA: {report['qa_status']}", size=11, bold=False)
    _add_paragraph(f"بازبینی اپراتور: {report['operator_review']}", size=11, bold=False)

    # ── محتوا ──────────────────────────────────────────────────────────────
    _add_heading("────── محتوا ──────")
    for label, value in report["content"]:
        _add_paragraph(f"{label}: {value}", size=11, bold=False)

    # ── فراداده ────────────────────────────────────────────────────────────
    _add_heading("────── فراداده ──────")
    for label, value in report["metadata"]:
        _add_paragraph(f"• {label}: {value}", size=10, bold=False)

    # ── منابع ──────────────────────────────────────────────────────────────
    _add_heading("────── منابع ──────")
    _add_paragraph("پیوست‌ها:", size=10, bold=False)
    for r in report["resources"]:
        _add_paragraph(f"  {r}", size=10, bold=False)

    # ── موارد حل‌نشده ──────────────────────────────────────────────────────
    _add_heading("────── موارد حل‌نشده ──────")
    for item in report["unresolved"]:
        _add_paragraph(f"• {item}", size=10, bold=False)

    # ── چک‌لیست نهایی اپراتور ──────────────────────────────────────────────
    _add_heading("────── چک‌لیست نهایی اپراتور ──────")
    for item in report["checklist"]:
        _add_paragraph(f"[ ] {item}", size=11, bold=False)

    _add_paragraph(report["footer"], size=8.5, bold=False, color="666666")

    doc.save(out_path)
    logger.info("DOCX دانش ساخته شد: %s", out_path)
    return out_path
